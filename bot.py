import os
import re
import asyncio
from datetime import datetime
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.utils.keyboard import InlineKeyboardBuilder
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document

# ИСПРАВЛЕНО: Вместо ручной вставки токена в код, облако будет брать его из безопасных переменных
BOT_TOKEN = os.environ.get("BOT_TOKEN")

USERS_DB = {}

def init_user(user_id):
    if user_id not in USERS_DB:
        USERS_DB[user_id] = {
            "balance": 3, "is_premium": False, "markup_mode": "split",
            "markups": {"default": 0.03, "40": 0.02, "50": 0.03, "60": 0.04, "70": 0.05, "80": 0.06, "90": 0.07}
        }

def parse_data(text):
    parsed = []
    pattern = re.compile(r"^\s*(\d+)\s*(hb|qb)\s+(.*?)\s+(\d+(?:-\d+)?)\s+(.+)$", re.IGNORECASE)
    box_counter = 1
    for line in text.strip().split("\n"):
        line = line.strip()
        if not line: continue
        match = pattern.match(line)
        if match:
            b_qty = int(match.group(1))
            b_type, var, cm_raw, pr_raw = match.group(2), match.group(3), match.group(4), match.group(5)
            prices = []
            for p in pr_raw.split():
                try: prices.append(float(p.replace(",", ".")))
                except ValueError: continue
            if not prices: prices = [0.0]
            if "-" in cm_raw and len(prices) >= 2:
                sizes = cm_raw.split("-")
                if len(sizes) >= 2:
                    parsed.append([box_counter, b_qty, b_type, var, sizes.strip(), prices])
                    parsed.append([box_counter, b_qty, b_type, var, sizes.strip(), prices])
                    box_counter += 1
                    continue
            parsed.append([box_counter, b_qty, b_type, var, str(cm_raw), prices])
            box_counter += 1
    return parsed
def generate_bot_documents(data_rows, user_settings, user_id):
    today_str = datetime.today().strftime("%d.%m.%Y")
    markups = user_settings["markups"]
    if user_settings["markup_mode"] == "single":
        v = markups["default"]
        act_m = {"default": v, "40": v, "50": v, "60": v, "70": v, "80": v, "90": v}
    else: act_m = markups

    f_title, f_header, f_data = Font(name="Arial", size=11, bold=True), Font(name="Arial", size=10, bold=True, color="FFFFFF"), Font(name="Arial", size=10)
    fill_h = PatternFill(start_color="1F497D", end_color="1F497D", fill_type="solid")
    border_t = Border(left=Side(style='thin', color='D9D9D9'), right=Side(style='thin', color='D9D9D9'), top=Side(style='thin', color='D9D9D9'), bottom=Side(style='thin', color='D9D9D9'))
    a_center, a_left = Alignment(horizontal="center", vertical="center"), Alignment(horizontal="left", vertical="center")

    p_ex1, p_ex2, p_doc, p_txt = f"Free_Sales_{user_id}.xlsx", f"SALES_OFFERT_{user_id}.xlsx", f"Free_Sales_{user_id}.docx", f"List_{user_id}.txt"

    wb1 = openpyxl.Workbook(); ws1 = wb1.active; ws1.title = "Sales"
    ws1["A1"], ws1["D1"] = "Free Sales Today", today_str
    ws1["A1"].font, ws1["D1"].font = f_title, f_title
    ws1.freeze_panes = "A4"
    headers1 = ["BX", "BOX", "BOX TYPE", "VARIETIES", "CM", "UNIT PRICE", "SALES PRICE"]
    for idx, h in enumerate(headers1, 1):
        c = ws1.cell(3, idx, h); c.font, c.fill, c.alignment = f_header, fill_h, a_center
    for i, row in enumerate(data_rows, 4):
        ws1.cell(i, 1, value=f"BX{row}").alignment = a_center
        for j in range(2, 7): ws1.cell(i, j, value=row[j-1]).alignment = a_left if j == 4 else a_center
        for j in range(1, 7): ws1.cell(i, j).font, ws1.cell(i, j).border = f_data, border_t
        ws1.cell(i, 7, f"=F{i}+{act_m.get(str(row), act_m['default'])}").font = f_data
        ws1.cell(i, 7).border, ws1.cell(i, 7).alignment = border_t, a_center
    wb1.save(p_ex1)

    wb2 = openpyxl.Workbook(); ws2 = wb2.active; ws2.title = "OFFERT"
    ws2["A1"], ws2["D1"] = "SALES OFFERT", today_str
    ws2["A1"].font, ws2["D1"].font = f_title, f_title
    ws2.freeze_panes = "A4"
    headers2 = ["BX", "BOX", "BOX TYPE", "VARIETIES", "CM", "SALES PRICE"]
    for idx, h in enumerate(headers2, 1):
        c = ws2.cell(3, idx, h); c.font, c.fill, c.alignment = f_header, fill_h, a_center
    for i, row in enumerate(data_rows, 4):
        ws2.cell(i, 1, value=f"BX{row}").alignment = a_center
        for j in range(2, 6): ws2.cell(i, j, value=row[j-1]).alignment = a_left if j == 4 else a_center
        for j in range(1, 6): ws2.cell(i, j).font, ws2.cell(i, j).border = f_data, border_t
        final_p = round(row + act_m.get(str(row), act_m["default"]), 2)
        c_sales = ws2.cell(i, 6, value=final_p); c_sales.font, c_sales.border, c_sales.alignment = f_data, border_t, a_center
    wb2.save(p_ex2)

    doc = Document(); doc.add_heading("Free Sales Today", level=1)
    table = doc.add_table(rows=1, cols=6); table.style = "Table Grid"
    for idx, h in enumerate(headers2): table.rows.cells[idx].text = h
    for row in data_rows:
        sales_p = round(row + act_m.get(str(row), act_m["default"]), 2)
        rc = table.add_row().cells
        for j, v in enumerate([f"BX{row}", str(row), str(row), str(row), str(row), str(sales_p)]): rc[j].text = v
    doc.save(p_doc)

    txt_content = ""
    with open(p_txt, "w", encoding="utf-8") as f:
        for row in data_rows:
            sales_p = round(row + act_m.get(str(row), act_m["default"]), 2)
            line = f"BX{row} {row} {row} {row} {row} {sales_p}\n"
            f.write(line); txt_content += line
    return p_ex1, p_ex2, p_doc, p_txt, txt_content

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

def get_main_menu(user_id):
    b = InlineKeyboardBuilder()
    b.button(text="⚙️ Настройка наценок", callback_data="settings_menu")
    b.button(text="📊 Мой баланс", callback_data="show_balance")
    return b.adjust(1).as_markup()

@dp.message(Command("start"))
async def start_cmd(message: types.Message):
    init_user(message.from_user.id)
    await message.answer("🌿 *Добро пожаловать!*\n\nОтправьте мне файл `.txt` с остатками, и я соберу полный пакет документов.", parse_mode="Markdown", reply_markup=get_main_menu(message.from_user.id))

@dp.callback_query(F.data == "show_balance")
async def show_balance(c: types.CallbackQuery):
    init_user(c.from_user.id)
    u = USERS_DB[c.from_user.id]
    await c.message.answer(f"📊 *Ваш аккаунт:*\nОстаток: {u['balance']} конвертаций", parse_mode="Markdown"); await c.answer()

@dp.callback_query(F.data == "settings_menu")
async def settings_menu(c: types.CallbackQuery):
    init_user(c.from_user.id)
    mode_text = "Единая" if USERS_DB[c.from_user.id]["markup_mode"] == "single" else "Раздельная по длинам (CM)"
    b = InlineKeyboardBuilder()
    b.button(text="🔄 Сменить режим", callback_data="switch_markup_mode")
    b.button(text="🔙 Назад", callback_data="main_menu")
    await c.message.edit_text(f"⚙️ *Настройка наценок*\nРежим: *{mode_text}*", parse_mode="Markdown", reply_markup=b.adjust(1).as_markup())

@dp.callback_query(F.data == "switch_markup_mode")
async def switch_markup_mode(c: types.CallbackQuery):
    init_user(c.from_user.id)
    USERS_DB[c.from_user.id]["markup_mode"] = "single" if USERS_DB[c.from_user.id]["markup_mode"] == "split" else "split"
    await settings_menu(c)

@dp.callback_query(F.data == "main_menu")
async def main_menu(c: types.CallbackQuery):
    await c.message.edit_text("🌿 Отправьте мне `.txt` файл для обработки.", reply_markup=get_main_menu(c.from_user.id))

@dp.message(F.document)
async def handle_document(message: types.Message):
    u_id = message.from_user.id
    init_user(u_id)
    if not message.document.file_name.endswith(".txt"):
        await message.answer("⚠️ Пожалуйста, присылайте файлы только в формате `.txt`."); return
    status_msg = await message.answer("⏳ Рассчитываю прайс-лист...")
    try:
        file_info = await bot.get_file(message.document.file_id)
        file_content = await bot.download_file(file_info.file_path)
        text_data = file_content.read().decode("utf-8")
        data_rows = parse_data(text_data)
        if not data_rows: await status_msg.edit_text("❌ Не удалось распознать структуру цветов."); return
        p_ex1, p_ex2, p_doc, p_txt, txt_content = generate_bot_documents(data_rows, USERS_DB[u_id], u_id)
        await message.reply_document(types.FSInputFile(p_ex1), caption="1. Свободная продажа сегодня (Excel)")
        await message.reply_document(types.FSInputFile(p_ex2), caption="2. SALES OFFERT (Excel)")
        await message.reply_document(types.FSInputFile(p_doc), caption="3. Таблица Свободная продажа (Word)")
        await message.answer(f"📋 *Текст для вставки в чаты:*\n\n`{txt_content[:3500]}`", parse_mode="Markdown")
        await status_msg.delete()
        for f in [p_ex1, p_ex2, p_doc, p_txt]:
            if os.path.exists(f): os.remove(f)
    except Exception as e: await status_msg.edit_text(f"❌ Ошибка: {str(e)}")

async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
