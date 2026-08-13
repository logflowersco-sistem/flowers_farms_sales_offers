import os
import re
from datetime import datetime
import flet as ft
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document

def parse_data(text):
    parsed = []
    pattern = re.compile(r"^\s*(\d+)\s*(hb|qb)\s+(.*?)\s+(\d+(?:-\d+)?)\s+(.+)$", re.IGNORECASE)
    box_counter = 1
    for line in text.strip().split("\n"):
        line = line.strip()
        if not line: continue
        match = pattern.match(line)
        if match:
            b_qty = int(match.group(1))
            b_type, var, cm_raw, pr_raw = match.group(2), match.group(3), match.group(4), match.group(5)
            prices = []
            for p in pr_raw.split():
                try: prices.append(float(p.replace(",", ".")))
                except ValueError: continue
            if not prices: prices = [0.0]
            
            if "-" in cm_raw and len(prices) >= 2:
                sizes = cm_raw.split("-")
                if len(sizes) >= 2:
                    s1 = sizes.strip()
                    s2 = sizes.strip()
                    p1 = prices
                    p2 = prices
                    parsed.append([box_counter, b_qty, b_type, var, s1, p1])
                    parsed.append([box_counter, b_qty, b_type, var, s2, p2])
                    box_counter += 1
                    continue
            
            single_price = prices
            parsed.append([box_counter, b_qty, b_type, var, str(cm_raw), single_price])
            box_counter += 1
    return parsed
def main(page: ft.Page):
    page.title = "Flower Converter v3.0"
    page.theme_mode = ft.ThemeMode.LIGHT
    page.scroll = ft.ScrollMode.AUTO
    page.padding = 20

    target_dir = "/storage/emulated/0/Download"

    input_text_field = ft.TextField(
        label="Paste your flower price text here",
        multiline=True,
        min_lines=5,
        max_lines=10,
        hint_text="Example:\n1 qb 3D 50 0.28\n2 hb Alba 60 0.35"
    )

    entry_m_def = ft.TextField(label="Default / Single Markup", value="0.03", width=200)
    entry_m_40 = ft.TextField(label="40 cm", value="0.02", width=120)
    entry_m_50 = ft.TextField(label="50 cm", value="0.03", width=120)
    entry_m_60 = ft.TextField(label="60 cm", value="0.04", width=120)
    entry_m_70 = ft.TextField(label="70 cm", value="0.05", width=120)
    entry_m_80 = ft.TextField(label="80 cm", value="0.06", width=120)
    entry_m_90 = ft.TextField(label="90 cm", value="0.07", width=120)

    def mode_changed(e):
        is_single = markup_mode.value == "single"
        for field in [entry_m_40, entry_m_50, entry_m_60, entry_m_70, entry_m_80, entry_m_90]:
            field.disabled = is_single
        page.update()

    markup_mode = ft.RadioGroup(
        content=ft.Column([
            ft.Radio(value="split", label="Split markups by size (CM)"),
            ft.Radio(value="single", label="Single markup for all sizes")
        ]),
        value="split",
        on_change=mode_changed
    )

    def run_conversion(e):
        selected_file_content = input_text_field.value.strip()
        if not selected_file_content:
            page.snack_bar = ft.SnackBar(ft.Text("Please paste text first!"))
            page.snack_bar.open = True
            page.update()
            return
            
        try:
            data_rows = parse_data(selected_file_content)
            if not data_rows:
                page.snack_bar = ft.SnackBar(ft.Text("Error parsing data! Check text format."))
                page.snack_bar.open = True
                page.update()
                return

            today_str = datetime.today().strftime("%d.%m.%Y")
            
            if markup_mode.value == "single":
                v = float(entry_m_def.value.replace(",", "."))
                markups = {"default": v, "40": v, "50": v, "60": v, "70": v, "80": v, "90": v}
            else:
                markups = {
                    "default": float(entry_m_def.value.replace(",", ".")), "40": float(entry_m_40.value.replace(",", ".")),
                    "50": float(entry_m_50.value.replace(",", ".")), "60": float(entry_m_60.value.replace(",", ".")),
                    "70": float(entry_m_70.value.replace(",", ".")), "80": float(entry_m_80.value.replace(",", ".")),
                    "90": float(entry_m_90.value.replace(",", "."))
                }

            f_title, f_header, f_data = Font(name="Arial", size=11, bold=True), Font(name="Arial", size=10, bold=True, color="FFFFFF"), Font(name="Arial", size=10)
            fill_h = PatternFill(start_color="1F497D", end_color="1F497D", fill_type="solid")
            border_t = Border(left=Side(style='thin', color='D9D9D9'), right=Side(style='thin', color='D9D9D9'), top=Side(style='thin', color='D9D9D9'), bottom=Side(style='thin', color='D9D9D9'))
            a_center, a_left = Alignment(horizontal="center", vertical="center"), Alignment(horizontal="left", vertical="center")

            wb1 = openpyxl.Workbook(); ws1 = wb1.active; ws1.title = "Sales"
            ws1["A1"], ws1["D1"] = "Free Sales Today", today_str
            ws1["A1"].font, ws1["D1"].font = f_title, f_title
            ws1.freeze_panes = "A4"
            headers1 = ["BX", "BOX", "BOX TYPE", "VARIETIES", "CM", "UNIT PRICE", "SALES PRICE"]
            for idx, h in enumerate(headers1, 1):
                c = ws1.cell(3, idx, h); c.font, c.fill, c.alignment = f_header, fill_h, a_center
            for i, row in enumerate(data_rows, 4):
                ws1.cell(i, 1, value=f"BX{row}").alignment = a_center
                for j in range(2, 7): ws1.cell(i, j, value=row[j-1]).alignment = a_left if j == 4 else a_center
                for j in range(1, 7): ws1.cell(i, j).font, ws1.cell(i, j).border = f_data, border_t
                cm_s = str(row)
                ws1.cell(i, 7, f"=F{i}+{markups.get(cm_s, markups['default'])}").font = f_data
                ws1.cell(i, 7).border, ws1.cell(i, 7).alignment = border_t, a_center
            for idx in range(1, 8):
                m_len = max(len(str(ws1.cell(r, idx).value or '')) for r in range(1, ws1.max_row + 1))
                ws1.column_dimensions[openpyxl.utils.get_column_letter(idx)].width = max(m_len + 4, 11)
            
            if not os.path.exists(target_dir): os.makedirs(target_dir, exist_ok=True)
            wb1.save(os.path.join(target_dir, "Free_Sales_Today.xlsx"))

            wb2 = openpyxl.Workbook(); ws2 = wb2.active; ws2.title = "OFFERT"
            ws2["A1"], ws2["D1"] = "SALES OFFERT", today_str
            ws2["A1"].font, ws2["D1"].font = f_title, f_title
            ws2.freeze_panes = "A4"
            headers2 = ["BX", "BOX", "BOX TYPE", "VARIETIES", "CM", "SALES PRICE"]
            for idx, h in enumerate(headers2, 1):
                c = ws2.cell(3, idx, h); c.font, c.fill, c.alignment = f_header, fill_h, a_center
            for i, row in enumerate(data_rows, 4):
                ws2.cell(i, 1, value=f"BX{row}").alignment = a_center
                for j in range(2, 6): ws2.cell(i, j, value=row[j-1]).alignment = a_left if j == 4 else a_center
                for j in range(1, 6): ws2.cell(i, j).font, ws2.cell(i, j).border = f_data, border_t
                cm_s = str(row)
                final_p = round(row + markups.get(cm_s, markups["default"]), 2)
                c_sales = ws2.cell(i, 6, value=final_p); c_sales.font, c_sales.border, c_sales.alignment = f_data, border_t, a_center
            for idx in range(1, 7):
                m_len = max(len(str(ws2.cell(r, idx).value or '')) for r in range(1, ws2.max_row + 1))
                ws2.column_dimensions[openpyxl.utils.get_column_letter(idx)].width = max(m_len + 4, 11)
            wb2.save(os.path.join(target_dir, f"SALES_OFFERT_{today_str}.xlsx"))

            doc = Document(); doc.add_heading("Free Sales Today", level=1); doc.add_paragraph(f"Date: {today_str}")
            table = doc.add_table(rows=1, cols=6); table.style = "Table Grid"
            hdr_cells = table.rows.cells
            for idx, h in enumerate(headers2): hdr_cells[idx].text = h
            for row in data_rows:
                cm_s = str(row)
                sales_p = round(row + markups.get(cm_s, markups["default"]), 2)
                row_cells = table.add_row().cells
                vals = [f"BX{row}", str(row), str(row), str(row), str(row), str(sales_p)]
                for j, val in enumerate(vals): row_cells[j].text = val
            doc.save(os.path.join(target_dir, "Free_Sales_Today.docx"))

            with open(os.path.join(target_dir, "Free_Sales_Today_List.txt"), "w", encoding="utf-8") as f:
                for row in data_rows:
                    cm_s = str(row)
                    sales_p = round(row + markups.get(cm_s, markups["default"]), 2)
                    f.write(f"BX{row} {row}{row} {row} {row} {sales_p}\n")

            page.dialog = ft.AlertDialog(
                title=ft.Text("Success!"),
                content=ft.Text("All 4 files successfully saved to 'Download' folder!"),
                actions=[ft.TextButton("OK", on_click=lambda x: page.close_dialog())]
            )
            page.dialog.open = True
        except Exception as ex:
            page.dialog = ft.AlertDialog(title=ft.Text("Error"), content=ft.Text(str(ex)))
            page.dialog.open = True
        page.update()

    btn_start = ft.Button("Generate Documents Packet", on_click=run_conversion, height=50)

    page.add(
        ft.Text("Flower Converter App v3.0", style=ft.TextThemeStyle.HEADLINE_SMALL, weight=ft.FontWeight.BOLD, color="blue"),
        ft.Divider(),
        input_text_field,
        ft.Container(height=10),
        ft.Text("Markup Mode:", weight=ft.FontWeight.BOLD),
        markup_mode,
        ft.Divider(),
        ft.Text("Markup Parameters:", weight=ft.FontWeight.BOLD),
        ft.Row([entry_m_def], alignment=ft.MainAxisAlignment.START),
        ft.Row([entry_m_40, entry_m_50, entry_m_60], wrap=True),
        ft.Row([entry_m_70, entry_m_80, entry_m_90], wrap=True),
        ft.Container(height=20),
        btn_start
    )

# ИСПРАВЛЕНО: Добавлен жесткий порт 8080 для веб-доступа
ft.run(main, view=ft.AppView.WEB_BROWSER, port=8080)
