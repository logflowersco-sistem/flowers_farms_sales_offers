import streamlit as st
import re
import os
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
import io

def parse_data(text):
    parsed = []
    pattern = re.compile(r"^\s*(\d+)\s*(hb|qb)\s+(.*?)\s+(\d+(?:-\d+)?)\s+(.+)$", re.IGNORECASE)
    box_counter = 1
    for line in text.strip().split("\n"):
        line = line.strip()
        if not line: continue
        match = pattern.match(line)
        if match:
            b_qty = int(match.group(1))
            b_type, var, cm_raw, pr_raw = match.group(2), match.group(3), match.group(4), match.group(5)
            prices = []
            for p in pr_raw.split():
                try: prices.append(float(p.replace(",", ".")))
                except ValueError: continue
            if not prices: prices = [0.0]
            
            if "-" in cm_raw and len(prices) >= 2:
                sizes = cm_raw.split("-")
                if len(sizes) >= 2:
                    parsed.append([box_counter, b_qty, b_type, var, sizes[0].strip(), prices[0]])
                    parsed.append([box_counter, b_qty, b_type, var, sizes[1].strip(), prices[1]])
                    box_counter += 1
                    continue
            
            parsed.append([box_counter, b_qty, b_type, var, str(cm_raw), prices[0]])
            box_counter += 1
    return parsed
st.set_page_config(page_title="Flower Converter v3.0", page_icon="🌿", layout="centered")

st.title("🌿 Flower Converter Pro")
st.write("Paste your raw price list text below to generate the standard business documents pack instantly.")

# Большое и удобное поле ввода для текста на смартфоне
input_text = st.text_area("Paste your price text here:", height=200, placeholder="Example:\n1 qb 3D 50 0.28\n2 hb Alba 60 0.35")

st.subheader("⚙️ Markup Settings")
mode = st.radio("Markup Mode:", ["Split markups by size (CM)", "Single markup for all sizes"])

col1, col2, col3 = st.columns(3)
with col1:
    m_def = st.number_input("Default / Single:", value=0.03, step=0.01, format="%.2f")
    m_40 = st.number_input("40 cm:", value=0.02, step=0.01, format="%.2f", disabled=(mode == "Single markup for all sizes"))
    m_50 = st.number_input("50 cm:", value=0.03, step=0.01, format="%.2f", disabled=(mode == "Single markup for all sizes"))
with col2:
    m_60 = st.number_input("60 cm:", value=0.04, step=0.01, format="%.2f", disabled=(mode == "Single markup for all sizes"))
    m_70 = st.number_input("70 cm:", value=0.05, step=0.01, format="%.2f", disabled=(mode == "Single markup for all sizes"))
with col3:
    m_80 = st.number_input("80 cm:", value=0.06, step=0.01, format="%.2f", disabled=(mode == "Single markup for all sizes"))
    m_90 = st.number_input("90 cm:", value=0.07, step=0.01, format="%.2f", disabled=(mode == "Single markup for all sizes"))

if st.button("🚀 Generate Documents Packet", type="primary", use_container_width=True):
    if not input_text.strip():
        st.error("Please paste price text first!")
    else:
        data_rows = parse_data(input_text)
        if not data_rows:
            st.error("Error parsing data! Check text format.")
        else:
            today_str = datetime.today().strftime("%d.%m.%Y")
            markups = {"default": m_def, "40": m_40, "50": m_50, "60": m_60, "70": m_70, "80": m_80, "90": m_90}
            if mode == "Single markup for all sizes":
                markups = {"default": m_def, "40": m_def, "50": m_def, "60": m_def, "70": m_def, "80": m_def, "90": m_def}

            # --- Логика генерации в память сервера ---
            f_title, f_header, f_data = Font(name="Arial", size=11, bold=True), Font(name="Arial", size=10, bold=True, color="FFFFFF"), Font(name="Arial", size=10)
            fill_h = PatternFill(start_color="1F497D", end_color="1F497D", fill_type="solid")
            border_t = Border(left=Side(style='thin', color='D9D9D9'), right=Side(style='thin', color='D9D9D9'), top=Side(style='thin', color='D9D9D9'), bottom=Side(style='thin', color='D9D9D9'))
            a_center, a_left = Alignment(horizontal="center", vertical="center"), Alignment(horizontal="left", vertical="center")

            # Excel 1
            wb1 = openpyxl.Workbook(); ws1 = wb1.active; ws1.title = "Sales"
            ws1["A1"], ws1["D1"] = "Free Sales Today", today_str
            ws1["A1"].font, ws1["D1"].font = f_title, f_title
            ws1.freeze_panes = "A4"
            headers1 = ["BX", "BOX", "BOX TYPE", "VARIETIES", "CM", "UNIT PRICE", "SALES PRICE"]
            for idx, h in enumerate(headers1, 1):
                c = ws1.cell(3, idx, h); c.font, c.fill, c.alignment = f_header, fill_h, a_center
            for i, row in enumerate(data_rows, 4):
                ws1.cell(i, 1, value=f"BX{row[0]}").alignment = a_center
                for j in range(2, 7): ws1.cell(i, j, value=row[j-1]).alignment = a_left if j == 4 else a_center
                for j in range(1, 7): ws1.cell(i, j).font, ws1.cell(i, j).border = f_data, border_t
                ws1.cell(i, 7, f"=F{i}+{markups.get(str(row[4]), markups['default'])}").font = f_data
                ws1.cell(i, 7).border, ws1.cell(i, 7).alignment = border_t, a_center
            for idx in range(1, 8):
                m_len = max(len(str(ws1.cell(r, idx).value or '')) for r in range(1, ws1.max_row + 1))
                ws1.column_dimensions[openpyxl.utils.get_column_letter(idx)].width = max(m_len + 4, 11)
            
            ex1_buffer = io.BytesIO()
            wb1.save(ex1_buffer)
            ex1_buffer.seek(0)

            # Excel 2
            wb2 = openpyxl.Workbook(); ws2 = wb2.active; ws2.title = "OFFERT"
            ws2["A1"], ws2["D1"] = "SALES OFFERT", today_str
            ws2["A1"].font, ws2["D1"].font = f_title, f_title
            ws2.freeze_panes = "A4"
            headers2 = ["BX", "BOX", "BOX TYPE", "VARIETIES", "CM", "SALES PRICE"]
            for idx, h in enumerate(headers2, 1):
                c = ws2.cell(3, idx, h); c.font, c.fill, c.alignment = f_header, fill_h, a_center
            for i, row in enumerate(data_rows, 4):
                ws2.cell(i, 1, value=f"BX{row[0]}").alignment = a_center
                for j in range(2, 6): ws2.cell(i, j, value=row[j-1]).alignment = a_left if j == 4 else a_center
                for j in range(1, 6): ws2.cell(i, j).font, ws2.cell(i, j).border = f_data, border_t
                final_p = round(row[5] + markups.get(str(row[4]), markups["default"]), 2)
                c_sales = ws2.cell(i, 6, value=final_p); c_sales.font, c_sales.border, c_sales.alignment = f_data, border_t, a_center
            for idx in range(1, 7):
                m_len = max(len(str(ws2.cell(r, idx).value or '')) for r in range(1, ws2.max_row + 1))
                ws2.column_dimensions[openpyxl.utils.get_column_letter(idx)].width = max(m_len + 4, 11)
            
            ex2_buffer = io.BytesIO()
            wb2.save(ex2_buffer)
            ex2_buffer.seek(0)

            # Word Doc
            doc = Document(); doc.add_heading("Free Sales Today", level=1); doc.add_paragraph(f"Date: {today_str}")
            table = doc.add_table(rows=1, cols=6); table.style = "Table Grid"
            for idx, h in enumerate(headers2): table.rows[0].cells[idx].text = h
            for row in data_rows:
                sales_p = round(row[5] + markups.get(str(row[4]), markups["default"]), 2)
                rc = table.add_row().cells
                vals = [f"BX{row[0]}", str(row[1]), str(row[2]), str(row[3]), str(row[4]), str(sales_p)]
                for j, val in enumerate(vals): rc[j].text = val
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Блокнот
            txt_content = ""
            for row in data_rows:
                sales_p = round(row[5] + markups.get(str(row[4]), markups["default"]), 2)
                txt_content += f"BX{row[0]} {row[1]}{row[2]} {row[3]} {row[4]} {sales_p}\n"

            st.success("🎉 Documents generated successfully! Click buttons below to download:")
            
            st.download_button("📥 Download Free Sales Today (Excel)", data=ex1_buffer, file_name="Free_Sales_Today.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
            st.download_button("📥 Download SALES OFFERT (Excel)", data=ex2_buffer, file_name=f"SALES_OFFERT_{today_str}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
            st.download_button("📥 Download Free Sales Today (Word)", data=doc_buffer, file_name="Free_Sales_Today.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            
            st.subheader("📋 Компактный текст для чатов WhatsApp / Telegram:")
            st.code(txt_content, language="text")
